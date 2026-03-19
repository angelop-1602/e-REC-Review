"""
Document generation functions for SPUP REC System
"""
import os
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Import config values - handle both relative and absolute imports
try:
    from .config import TEMPLATE_LETTER
    from .utils import generate_output_filename, format_date
except ImportError:
    from config import TEMPLATE_LETTER
    from utils import generate_output_filename, format_date


def create_word_table(doc, data, headers, title):
    """Create a properly formatted table in Word document"""
    if not data:
        return
    
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Inches(0.16)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    if len(headers) > 5:
        # For complete list table - adjust widths based on number of columns
        if len(headers) == 8:  # With Research Title
            widths = [Inches(0.8), Inches(1.5), Inches(1.8), Inches(2.0), Inches(1.2), Inches(1.2), Inches(1.2), Inches(1.2)]
        else:  # Without Research Title (7 columns)
            widths = [Inches(0.8), Inches(1.5), Inches(2.0), Inches(1.2), Inches(1.2), Inches(1.2), Inches(1.2)]
    else:
        # For summary table (3 columns)
        widths = [Inches(2.5), Inches(1.5), Inches(1.8)]
    
    for i, width in enumerate(widths[:len(headers)]):
        for cell in table.columns[i].cells:
            cell.width = width
    
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Inches(0.12)
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, header in enumerate(headers):
            cell_value = str(row_data.get(header, ''))
            row_cells[i].text = cell_value
            row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Inches(0.10)
    
    doc.add_paragraph()


def create_word_document(template_path, output_path, replacements, data, summary_data):
    """Create Word document with properly formatted tables"""
    try:
        doc = Document(template_path)
        
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    if key not in ['<<Complete_List_Table>>', '<<Summary_Table>>']:
                        paragraph.text = paragraph.text.replace(key, str(value))
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            if key not in ['<<Complete_List_Table>>', '<<Summary_Table>>']:
                                cell.text = cell.text.replace(key, str(value))
        
        for paragraph in doc.paragraphs:
            if '<<Complete_List_Table>>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<<Complete_List_Table>>', '')
                # Build headers from actual data columns - follow CSV order but move OR to first
                if data and len(data) > 0:
                    # Get columns in the order they appear in the CSV file
                    complete_headers = list(data[0].keys())
                    
                    # Find OR column (case-insensitive, handle variations)
                    or_column = None
                    for col in complete_headers:
                        if col.strip().upper() == 'OR':
                            or_column = col
                            break
                    
                    # Move OR to first position if it exists
                    if or_column and or_column in complete_headers:
                        complete_headers.remove(or_column)
                        complete_headers.insert(0, or_column)
                else:
                    # Fallback if no data
                    complete_headers = ['OR', 'SPUP REC Code', 'Principal Investigator', 'Course/Program', 'Reviewer #1', 'Reviewer #2', 'Reviewer #3']
                create_word_table(doc, data, complete_headers, "Complete List of Applications")
            
            if '<<Summary_Table>>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<<Summary_Table>>', '')
                summary_headers = ['Name of Reviewers', 'Number of Required Proposals', 'Honorarium (₱500 Per Proposal)']
                create_word_table(doc, summary_data, summary_headers, "Summary of Reviewers")
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '<<Complete_List_Table>>' in cell.text:
                        cell.text = cell.text.replace('<<Complete_List_Table>>', '')
                        # Build headers from actual data columns - follow CSV order but move OR to first
                        if data and len(data) > 0:
                            # Get columns in the order they appear in the CSV file
                            complete_headers = list(data[0].keys())
                            
                            # Find OR column (case-insensitive, handle variations)
                            or_column = None
                            for col in complete_headers:
                                if col.strip().upper() == 'OR':
                                    or_column = col
                                    break
                            
                            # Move OR to first position if it exists
                            if or_column and or_column in complete_headers:
                                complete_headers.remove(or_column)
                                complete_headers.insert(0, or_column)
                        else:
                            # Fallback if no data
                            complete_headers = ['OR', 'SPUP REC Code', 'Principal Investigator', 'Course/Program', 'Reviewer #1', 'Reviewer #2', 'Reviewer #3']
                        create_word_table(doc, data, complete_headers, "Complete List of Applications")
                        
                    if '<<Summary_Table>>' in cell.text:
                        cell.text = cell.text.replace('<<Summary_Table>>', '')
                        summary_headers = ['Name of Reviewers', 'Number of Required Proposals', 'Honorarium (₱500 Per Proposal)']
                        create_word_table(doc, summary_data, summary_headers, "Summary of Reviewers")
        
        doc.save(output_path)
        return True, "Document created successfully"
    except Exception as e:
        return False, f"Error creating document: {str(e)}"


def create_letter_document(date_today, level, month_year_display, amount, data, summary_data):
    """Create letter document with all replacements"""
    template_path = TEMPLATE_LETTER
    if not os.path.exists(template_path):
        return False, None, f"Template file not found: {template_path}"
    
    output_path = generate_output_filename("SPUP_REC_Letter")
    
    replacements = {
        "<<Date_Today>>": format_date(date_today, "%B %d, %Y"),
        "<<Level>>": level,
        "<<Month_Year>>": month_year_display,
        "<<Amount>>": str(amount)
    }
    
    success, message = create_word_document(template_path, output_path, replacements, data, summary_data)
    
    if success:
        return True, output_path, message
    else:
        return False, None, message

