"""
Voucher generation functions for SPUP REC System
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Import config values - handle both relative and absolute imports
try:
    from .config import TEMPLATE_VOUCHER, REVIEWERS_LIST
    from .utils import format_currency, generate_output_filename
except ImportError:
    from config import TEMPLATE_VOUCHER, REVIEWERS_LIST
    from utils import format_currency, generate_output_filename


def create_voucher_document(data, month_year_display, amount_per_review):
    """Create voucher document for all reviewers"""
    try:
        # Check if voucher template exists
        if not os.path.exists(TEMPLATE_VOUCHER):
            # Create a basic template if it doesn't exist
            doc = Document()
            doc.add_paragraph("SPUP REC Voucher Template")
            doc.save(TEMPLATE_VOUCHER)
        
        doc = Document(TEMPLATE_VOUCHER)
        
        # Collect reviewers and their applications
        reviewer_applications = {}
        
        for row in data:
            for col in ['Reviewer #1', 'Reviewer #2', 'Reviewer #3']:
                reviewer = row.get(col, '').strip()
                if reviewer and reviewer in REVIEWERS_LIST:
                    if reviewer not in reviewer_applications:
                        reviewer_applications[reviewer] = []
                    app_code = row.get('SPUP REC Code', row.get('OR', 'Unknown'))
                    reviewer_applications[reviewer].append(app_code)
        
        # Clear existing content
        doc.paragraphs.clear()
        
        first_reviewer = True
        
        # Create voucher for each reviewer in order
        for reviewer in REVIEWERS_LIST:
            if reviewer in reviewer_applications:
                applications = reviewer_applications[reviewer]
                
                # Add page break except for first reviewer
                if not first_reviewer:
                    doc.add_page_break()
                first_reviewer = False
                
                # Add header
                header_paragraph = doc.add_paragraph()
                header_run = header_paragraph.add_run(f"VOUCHER FOR: {reviewer}")
                header_run.bold = True
                header_run.font.size = Pt(14)
                header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                doc.add_paragraph()
                
                # Add period
                period_paragraph = doc.add_paragraph()
                period_run = period_paragraph.add_run(f"Period: {month_year_display}")
                period_run.font.size = Pt(12)
                period_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                doc.add_paragraph()
                
                # Create table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                table.autofit = True
                
                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Application Code'
                hdr_cells[1].text = 'Amount'
                
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(12)
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Add application rows
                for app_code in applications:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(app_code)
                    row_cells[1].text = format_currency(amount_per_review)
                    
                    row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                
                # Add total row
                total_row = table.add_row().cells
                total_row[0].text = "TOTAL"
                total_amount = len(applications) * amount_per_review
                total_row[1].text = format_currency(total_amount)
                
                for cell in total_row:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(12)
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Add signature fields
                doc.add_paragraph()
                doc.add_paragraph()
                
                date_paragraph = doc.add_paragraph("Date: ___________________")
                date_paragraph.runs[0].font.size = Pt(12)
                date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                received_paragraph = doc.add_paragraph("Received By: ___________________")
                received_paragraph.runs[0].font.size = Pt(12)
                received_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Generate filename
        output_filename = generate_output_filename("All_Vouchers")
        
        doc.save(output_filename)
        return True, output_filename, "Vouchers created successfully"
    except Exception as e:
        return False, None, f"Error creating vouchers: {str(e)}"

